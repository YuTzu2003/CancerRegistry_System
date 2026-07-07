import os
import io
import uuid
import zipfile
import base64
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.sync_api import sync_playwright

def generate_export_files(format_pdf, format_word, charts_data, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body { font-family: sans-serif; }
            .chart-section { margin-bottom: 40px; page-break-after: always; }
            .annual-report-table { 
                border-collapse: collapse; width: 100%; font-size: 10px; margin-bottom: 20px;
            }
            .annual-report-table th, .annual-report-table td { 
                border: 1px solid #ccc; padding: 4px;
            }
            .annual-report-table caption { font-weight: bold; font-size: 14px; margin-bottom: 10px; }
            .chart-img { max-width: 100%; height: auto; margin-bottom: 15px; display: block; }
            .llm-text { background: #f8f9fa; padding: 15px; border-radius: 5px; font-size: 12px; white-space: pre-wrap; margin-top: 15px; }
            .text-center { text-align: center !important; }
            .text-start { text-align: left !important; }
            .text-end { text-align: right !important; }
            .fw-bold { font-weight: bold !important; }
            .ps-4 { padding-left: 1.5rem !important; }
            .table-light { background-color: #f8f9fa !important; }
            .table-secondary { background-color: #e2e3e5 !important; }
            @page { size: A4 landscape; margin: 1cm; }
        </style>
    </head>
    <body>
    """
    
    image_bytes_map = {}
    
    for idx, chart in enumerate(charts_data):
        html_content += f'<div class="chart-section" id="section-{idx}">'
        html_content += f'<h2>{chart.get("title", "")}</h2>'
        
        # Table
        if chart.get('includeTable', True):
            table_html = chart.get('tableHtml', '')
            if table_html:
                html_content += table_html
                
        # Image
        if chart.get('includeChart', True):
            b64_image = chart.get('chartImage', '')
            if b64_image and b64_image.startswith('data:image'):
                html_content += f'<div class="chart-img-wrapper"><img src="{b64_image}" class="chart-img" data-idx="{idx}" /></div>'
                header, encoded = b64_image.split(",", 1)
                image_bytes_map[str(idx)] = base64.b64decode(encoded)
                
        # LLM text
        if chart.get('includeAi', True):
            llm_text = chart.get('llmText', '')
            if llm_text:
                html_content += f'<div class="llm-text"><strong>AI 分析敘述:</strong><br/>{llm_text}</div>'
        html_content += '</div>'    
    html_content += "</body></html>"
    
    pdf_bytes = None
    docx_bytes = None
    
    # 1. Generate PDF
    if format_pdf:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.set_content(html_content, wait_until="networkidle")
            page.emulate_media(media="print")
            pdf_bytes = page.pdf(format="A4", landscape=True, print_background=True)
            browser.close()
            
    # 2. Generate Word
    if format_word:
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
                
            doc = Document()
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            sections = soup.find_all('div', class_='chart-section')
            for sec_idx, sec in enumerate(sections):
                if sec_idx > 0:
                    doc.add_page_break()
                    
                title = sec.find('h2')
                if title:
                    doc.add_heading(title.text.strip(), level=1)
                    
                table_node = sec.find('table', class_='annual-report-table')
                if table_node:
                    caption = table_node.find('caption')
                    if caption:
                        doc.add_paragraph(caption.get_text(separator='\n').strip())
                        
                    rows = table_node.find_all('tr')
                    if rows:
                        # Determine exact grid dimensions
                        grid = {}
                        max_cols = 0
                        for r_idx, row in enumerate(rows):
                            cells = row.find_all(['th', 'td'])
                            c_idx = 0
                            for cell in cells:
                                # Find first empty spot in this row
                                while (r_idx, c_idx) in grid:
                                    c_idx += 1
                                
                                rowspan = int(cell.get('rowspan', 1))
                                colspan = int(cell.get('colspan', 1))
                                
                                # parse classes and styles
                                cell_classes = cell.get('class', [])
                                if isinstance(cell_classes, str): cell_classes = cell_classes.split()
                                row_classes = row.get('class', [])
                                if isinstance(row_classes, str): row_classes = row_classes.split()
                                table_classes = table_node.get('class', [])
                                if isinstance(table_classes, str): table_classes = table_classes.split()
                                
                                all_classes = table_classes + row_classes + cell_classes
                                
                                style_str = (cell.get('style', '') or '') + ';' + (row.get('style', '') or '')
                                is_bold = cell.name == 'th' or 'fw-bold' in all_classes or 'bold' in style_str or '900' in style_str
                                
                                alignment = WD_ALIGN_PARAGRAPH.CENTER
                                if 'text-start' in all_classes or 'ps-4' in cell_classes: alignment = WD_ALIGN_PARAGRAPH.LEFT
                                if 'text-center' in cell_classes: alignment = WD_ALIGN_PARAGRAPH.CENTER
                                elif 'text-end' in cell_classes: alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                
                                font_color = None
                                import re
                                color_match = re.search(r'color:\s*([^;]+)', style_str)
                                if color_match:
                                    font_color = color_match.group(1).strip()
                                
                                # Mark the grid
                                for r in range(rowspan):
                                    for c in range(colspan):
                                        grid[(r_idx + r, c_idx + c)] = {
                                            'text': cell.text.strip(),
                                            'is_head': is_bold,
                                            'alignment': alignment,
                                            'font_color': font_color,
                                            'main_r': r_idx,
                                            'main_c': c_idx,
                                            'rowspan': rowspan,
                                            'colspan': colspan
                                        }
                                
                                c_idx += colspan
                                max_cols = max(max_cols, c_idx)
                                
                        if max_cols > 0:
                            w_table = doc.add_table(rows=len(rows), cols=max_cols)
                            w_table.style = 'Table Grid'
                            
                            # Apply merges
                            merged_cells = set()
                            for (r, c), data in grid.items():
                                if (r, c) in merged_cells:
                                    continue
                                if data['rowspan'] > 1 or data['colspan'] > 1:
                                    main_cell = w_table.cell(data['main_r'], data['main_c'])
                                    target_cell = w_table.cell(data['main_r'] + data['rowspan'] - 1, data['main_c'] + data['colspan'] - 1)
                                    main_cell.merge(target_cell)
                                    for mr in range(data['rowspan']):
                                        for mc in range(data['colspan']):
                                            merged_cells.add((data['main_r'] + mr, data['main_c'] + mc))
                                    
                            # Write text and apply styles to the main cells
                            written = set()
                            for (r, c), data in grid.items():
                                mr, mc = data['main_r'], data['main_c']
                                if (mr, mc) not in written:
                                    w_cell = w_table.cell(mr, mc)
                                    w_cell.text = data['text']
                                    for paragraph in w_cell.paragraphs:
                                        paragraph.alignment = data['alignment']
                                        for run in paragraph.runs:
                                            run.font.size = Pt(8)
                                            if data['is_head']:
                                                run.font.bold = True
                                            if data['font_color']:
                                                if 'darkred' in data['font_color']: run.font.color.rgb = RGBColor(139, 0, 0)
                                                elif 'red' in data['font_color']: run.font.color.rgb = RGBColor(255, 0, 0)
                                    # Ensure minimal padding to prevent overlap
                                    written.add((mr, mc))
                                            
                img_node = sec.find('img', class_='chart-img')
                if img_node and img_node.get('data-idx'):
                    idx_str = img_node.get('data-idx')
                    if idx_str in image_bytes_map:
                        doc.add_paragraph()
                        img_stream = io.BytesIO(image_bytes_map[idx_str])
                        doc.add_picture(img_stream, width=Inches(9.0))
                        
                llm_node = sec.find('div', class_='llm-text')
                if llm_node:
                    doc.add_paragraph()
                    # Split by newlines so Word formats paragraphs properly
                    lines = llm_node.text.strip().split('\n')
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
        except Exception as e:
            print(f"Word export failed: {e}")
            
    # Return bytes in io.BytesIO
    if format_pdf and format_word:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zf:
            if pdf_bytes:
                zf.writestr("export_report.pdf", pdf_bytes)
            if docx_bytes:
                zf.writestr("export_report.docx", docx_bytes)
        zip_buffer.seek(0)
        return zip_buffer, "application/zip", "export_report.zip"
    elif format_pdf and pdf_bytes:
        return io.BytesIO(pdf_bytes), "application/pdf", "export_report.pdf"
    elif format_word and docx_bytes:
        return io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "export_report.docx"
    return None, None, None