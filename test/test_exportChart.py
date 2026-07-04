import os
import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from playwright.sync_api import sync_playwright

def convert_to_pdf_and_word(html_path, output_docx, output_pdf):
    html_abs_path = "file://" + os.path.abspath(html_path).replace("\\", "/")
    pdf_abs_path = os.path.abspath(output_pdf)
    docx_abs_path = os.path.abspath(output_docx)
    chart_img_path = os.path.abspath("temp_chart.png")

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(html_abs_path, wait_until="networkidle")
        page.wait_for_timeout(2000)
        
        canvas_locator = page.locator('#testChart canvas')
        if canvas_locator.count() > 0:
            canvas_locator.first.screenshot(path=chart_img_path)
            
        page.add_style_tag(content="""
            @page { size: A4 landscape; margin: 1cm; }
            .annual-report-table { 
                min-width: 0 !important; 
                width: 100% !important; 
                font-size: 10px !important; 
                table-layout: fixed;
            }
            .annual-report-table th, .annual-report-table td { 
                padding: 4px !important; 
                white-space: normal !important; 
                word-wrap: break-word !important;
            }
            .annual-report-table-wrap { 
                page-break-after: always; 
                margin-bottom: 20px;
            }""")
        
        page.emulate_media(media="print")
        page.pdf(path=pdf_abs_path, format="A4", landscape=True, print_background=True)
        print(f"Success! PDF 匯出完成: {pdf_abs_path}")
        
        browser.close()

    try:
        with open(os.path.abspath(html_path), 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
            
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        table_node = soup.find('table', class_='annual-report-table')
        if table_node:
            caption = table_node.find('caption')
            if caption:
                doc.add_paragraph(caption.text.strip())
                
            rows = table_node.find_all('tr')
            if rows:
                num_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                w_table = doc.add_table(rows=len(rows), cols=num_cols)
                w_table.style = 'Table Grid'
                
                for r_idx, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for c_idx, cell in enumerate(cells):
                        w_cell = w_table.cell(r_idx, c_idx)
                        w_cell.text = cell.text.strip()
                        for paragraph in w_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                if cell.name == 'th' or 'fw-bold' in cell.get('class', []):
                                    run.font.bold = True
        
        doc.add_page_break()
        
        chart_caption = soup.find('div', id='annualSexAgeChartCaption')
        if chart_caption:
            doc.add_paragraph(chart_caption.text.strip())
            
        if os.path.exists(chart_img_path):
            doc.add_picture(chart_img_path, width=Inches(9.5))
            
        doc.save(docx_abs_path)
        print(f"Success! Word 原生檔案建立完成: {docx_abs_path}")
        
    except Exception as e:
        print(f"Word 匯出失敗: {e}")
    finally:
        if os.path.exists(chart_img_path):
            os.remove(chart_img_path)

if __name__ == "__main__":
    convert_to_pdf_and_word("test_chart.html", "output.docx", "output.pdf")
